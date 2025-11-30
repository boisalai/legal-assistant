"""
Service d'extraction de texte multi-format.

Supporte l'extraction de texte a partir de:
- PDF (via pypdf, Docling, ou OCR)
- Word (.doc, .docx)
- Texte (.txt, .rtf)
- Markdown (.md)
- Audio (via transcription Whisper)
"""

import logging
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
import mimetypes

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Resultat de l'extraction de texte."""
    success: bool
    text: str = ""
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None
    extraction_method: str = "unknown"
    is_transcription: bool = False  # True pour les fichiers audio


# Types MIME supportes par categorie
SUPPORTED_TYPES = {
    # Documents PDF
    "application/pdf": "pdf",
    # Documents Word
    "application/msword": "word",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "word",
    # Documents texte
    "text/plain": "text",
    "text/rtf": "text",
    "application/rtf": "text",
    "text/markdown": "markdown",
    "text/x-markdown": "markdown",
    # Images (pour OCR)
    "image/jpeg": "image",
    "image/png": "image",
    "image/gif": "image",
    "image/webp": "image",
    "image/tiff": "image",
    # Audio (pour transcription)
    "audio/mpeg": "audio",
    "audio/mp3": "audio",
    "audio/mp4": "audio",
    "audio/x-m4a": "audio",
    "audio/wav": "audio",
    "audio/x-wav": "audio",
    "audio/webm": "audio",
    "audio/ogg": "audio",
    "audio/flac": "audio",
    "audio/aac": "audio",
    "audio/opus": "audio",
    "video/mp4": "audio",  # Peut contenir uniquement de l'audio
    "video/webm": "audio",
}

# Extensions de fichiers supportees
SUPPORTED_EXTENSIONS = {
    # PDF
    ".pdf": "pdf",
    # Word
    ".doc": "word",
    ".docx": "word",
    # Texte
    ".txt": "text",
    ".rtf": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    # Images
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".gif": "image",
    ".webp": "image",
    ".tiff": "image",
    ".tif": "image",
    # Audio
    ".mp3": "audio",
    ".mp4": "audio",
    ".m4a": "audio",
    ".wav": "audio",
    ".webm": "audio",
    ".ogg": "audio",
    ".opus": "audio",
    ".flac": "audio",
    ".aac": "audio",
    ".mp2": "audio",
    ".pcm": "audio",
    ".wma": "audio",
}


def get_file_category(filename: str, content_type: Optional[str] = None) -> Optional[str]:
    """
    Determine la categorie d'un fichier.

    Args:
        filename: Nom du fichier
        content_type: Type MIME (optionnel)

    Returns:
        Categorie ('pdf', 'word', 'text', 'markdown', 'image', 'audio') ou None
    """
    # D'abord essayer par extension
    ext = Path(filename).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[ext]

    # Puis par type MIME
    if content_type and content_type in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[content_type]

    # Essayer de deviner le type MIME
    guessed_type, _ = mimetypes.guess_type(filename)
    if guessed_type and guessed_type in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[guessed_type]

    return None


def is_supported_file(filename: str, content_type: Optional[str] = None) -> bool:
    """Verifie si un fichier est supporte."""
    return get_file_category(filename, content_type) is not None


class DocumentExtractionService:
    """
    Service d'extraction de texte multi-format.

    Methodes d'extraction par type:
    - PDF, Word, PowerPoint, Excel, images: MarkItDown (recommande)
    - PDF: pypdf (fallback rapide), Docling (avance), OCR (scans)
    - Word: python-docx (fallback)
    - Texte: lecture directe
    - Audio: Whisper (transcription)
    """

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self):
        """Verifie les dependances disponibles."""
        self._has_markitdown = False
        self._has_pypdf = False
        self._has_docx = False
        self._has_docling = False
        self._has_whisper = False
        self._has_pillow = False

        try:
            from markitdown import MarkItDown
            self._has_markitdown = True
            self._markitdown = MarkItDown()
            logger.info("MarkItDown disponible - utilisation pour l'extraction de documents")
        except ImportError:
            logger.warning("markitdown non installe - utilisation des methodes alternatives")

        try:
            import pypdf
            self._has_pypdf = True
        except ImportError:
            logger.warning("pypdf non installe - extraction PDF limitee")

        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("python-docx non installe - extraction Word limitee")

        try:
            from docling.document_converter import DocumentConverter
            self._has_docling = True
        except ImportError:
            logger.debug("Docling non installe - extraction avancee non disponible")

        try:
            import whisper
            self._has_whisper = True
        except ImportError:
            logger.debug("openai-whisper non installe - transcription audio limitee")

        try:
            from PIL import Image
            self._has_pillow = True
        except ImportError:
            logger.debug("Pillow non installe - extraction d'images limitee")

    async def extract(
        self,
        file_path: str | Path,
        content_type: Optional[str] = None,
        extraction_method: str = "auto",
        use_ocr: bool = False,
        language: str = "fr",
    ) -> ExtractionResult:
        """
        Extrait le texte d'un fichier.

        Args:
            file_path: Chemin vers le fichier
            content_type: Type MIME (optionnel)
            extraction_method: Methode d'extraction ('auto', 'pypdf', 'docling', 'ocr')
            use_ocr: Forcer l'utilisation de l'OCR
            language: Langue pour la transcription audio

        Returns:
            ExtractionResult avec le texte extrait
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return ExtractionResult(
                success=False,
                error=f"Fichier non trouve: {file_path}"
            )

        # Determiner la categorie du fichier
        category = get_file_category(file_path.name, content_type)

        if not category:
            return ExtractionResult(
                success=False,
                error=f"Type de fichier non supporte: {file_path.suffix}"
            )

        # Router vers la methode d'extraction appropriee
        if category == "pdf":
            return await self._extract_pdf(file_path, extraction_method, use_ocr)
        elif category == "word":
            return await self._extract_word(file_path)
        elif category == "text":
            return await self._extract_text(file_path)
        elif category == "markdown":
            return await self._extract_markdown(file_path)
        elif category == "image":
            return await self._extract_image(file_path)
        elif category == "audio":
            return await self._transcribe_audio(file_path, language)
        else:
            return ExtractionResult(
                success=False,
                error=f"Categorie non geree: {category}"
            )

    async def _extract_pdf(
        self,
        file_path: Path,
        method: str = "auto",
        use_ocr: bool = False,
    ) -> ExtractionResult:
        """Extrait le texte d'un PDF."""
        # Si OCR demande et Docling disponible
        if use_ocr and self._has_docling:
            return await self._extract_pdf_docling(file_path, use_vlm=True)

        # Si methode specifique demandee
        if method == "docling" and self._has_docling:
            return await self._extract_pdf_docling(file_path)
        elif method == "docling-vlm" and self._has_docling:
            return await self._extract_pdf_docling(file_path, use_vlm=True)

        # Par defaut: MarkItDown (meilleure qualite)
        if self._has_markitdown:
            return await self._extract_with_markitdown(file_path)

        # Fallback: pypdf (rapide mais moins bon)
        if self._has_pypdf:
            return await self._extract_pdf_pypdf(file_path)

        return ExtractionResult(
            success=False,
            error="Aucune bibliotheque PDF disponible (installer markitdown, pypdf ou docling)"
        )

    async def _extract_with_markitdown(self, file_path: Path) -> ExtractionResult:
        """Extraction avec MarkItDown (PDF, Word, PowerPoint, Excel, images)."""
        try:
            result = self._markitdown.convert(str(file_path))

            if not result or not result.text_content:
                return ExtractionResult(
                    success=False,
                    error="MarkItDown n'a pas pu extraire de texte",
                    extraction_method="markitdown"
                )

            return ExtractionResult(
                success=True,
                text=result.text_content,
                metadata={
                    "source": str(file_path),
                    "title": getattr(result, "title", None),
                },
                extraction_method="markitdown"
            )
        except Exception as e:
            logger.error(f"Erreur extraction MarkItDown: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="markitdown"
            )

    async def _extract_pdf_pypdf(self, file_path: Path) -> ExtractionResult:
        """Extraction PDF avec pypdf."""
        try:
            import pypdf

            reader = pypdf.PdfReader(str(file_path))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)

            return ExtractionResult(
                success=True,
                text=full_text,
                metadata={
                    "num_pages": len(reader.pages),
                    "has_text": bool(full_text.strip()),
                },
                extraction_method="pypdf"
            )
        except Exception as e:
            logger.error(f"Erreur extraction pypdf: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="pypdf"
            )

    async def _extract_pdf_docling(
        self,
        file_path: Path,
        use_vlm: bool = False,
    ) -> ExtractionResult:
        """Extraction PDF avec Docling."""
        try:
            from services.docling_service import get_docling_service

            service = get_docling_service(use_vlm=use_vlm)
            result = await service.extract_pdf(file_path)

            return ExtractionResult(
                success=result.success,
                text=result.text,
                metadata=result.metadata,
                error=result.error,
                extraction_method=result.extraction_method
            )
        except Exception as e:
            logger.error(f"Erreur extraction Docling: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="docling"
            )

    async def _extract_word(self, file_path: Path) -> ExtractionResult:
        """Extrait le texte d'un document Word."""
        # Utiliser MarkItDown si disponible (meilleure qualite)
        if self._has_markitdown:
            return await self._extract_with_markitdown(file_path)

        # Fallback: python-docx
        if not self._has_docx:
            return ExtractionResult(
                success=False,
                error="Aucune bibliotheque Word disponible (installer markitdown ou python-docx)",
                extraction_method="docx"
            )

        try:
            import docx

            doc = docx.Document(str(file_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extraire aussi les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)

            return ExtractionResult(
                success=True,
                text=full_text,
                metadata={
                    "num_paragraphs": len(doc.paragraphs),
                    "num_tables": len(doc.tables),
                },
                extraction_method="docx"
            )
        except Exception as e:
            logger.error(f"Erreur extraction Word: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="docx"
            )

    async def _extract_text(self, file_path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier texte brut."""
        try:
            # Essayer plusieurs encodages
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    text = file_path.read_text(encoding=encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                return ExtractionResult(
                    success=False,
                    error="Impossible de decoder le fichier texte",
                    extraction_method="text"
                )

            return ExtractionResult(
                success=True,
                text=text,
                metadata={
                    "encoding": used_encoding,
                    "size_chars": len(text),
                },
                extraction_method="text"
            )
        except Exception as e:
            logger.error(f"Erreur extraction texte: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="text"
            )

    async def _extract_markdown(self, file_path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier Markdown."""
        result = await self._extract_text(file_path)
        if result.success:
            result.extraction_method = "markdown"
        return result

    async def _extract_image(self, file_path: Path) -> ExtractionResult:
        """Extrait le texte d'une image via OCR."""
        if not self._has_docling:
            return ExtractionResult(
                success=False,
                error="Docling non installe pour l'OCR d'images",
                extraction_method="ocr"
            )

        try:
            from services.docling_service import get_docling_service

            service = get_docling_service(use_vlm=True)

            return ExtractionResult(
                success=False,
                error="L'extraction OCR d'images n'est pas encore implementee",
                extraction_method="ocr"
            )
        except Exception as e:
            logger.error(f"Erreur OCR image: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="ocr"
            )

    async def _transcribe_audio(
        self,
        file_path: Path,
        language: str = "fr",
    ) -> ExtractionResult:
        """Transcrit un fichier audio en texte."""
        if not self._has_whisper:
            # Fallback: utiliser l'API Whisper d'OpenAI ou autre service
            return await self._transcribe_audio_api(file_path, language)

        try:
            import whisper

            # Charger le modele (utiliser "base" pour la rapidite, "large" pour la qualite)
            model = whisper.load_model("base")

            # Transcrire
            result = model.transcribe(
                str(file_path),
                language=language if language != "auto" else None,
            )

            return ExtractionResult(
                success=True,
                text=result["text"],
                metadata={
                    "language": result.get("language", language),
                    "duration_seconds": result.get("duration", 0),
                },
                extraction_method="whisper",
                is_transcription=True
            )
        except Exception as e:
            logger.error(f"Erreur transcription Whisper: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                extraction_method="whisper",
                is_transcription=True
            )

    async def _transcribe_audio_api(
        self,
        file_path: Path,
        language: str = "fr",
    ) -> ExtractionResult:
        """Transcrit via une API externe (fallback)."""
        return ExtractionResult(
            success=False,
            error="Transcription audio non disponible. Installer openai-whisper: pip install openai-whisper",
            extraction_method="api",
            is_transcription=True
        )


# Singleton
_extraction_service: Optional[DocumentExtractionService] = None


def get_extraction_service() -> DocumentExtractionService:
    """Obtient l'instance singleton du service d'extraction."""
    global _extraction_service
    if _extraction_service is None:
        _extraction_service = DocumentExtractionService()
    return _extraction_service
